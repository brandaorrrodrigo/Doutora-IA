#!/usr/bin/env python3
"""
Generate DOCX templates for legal documents using python-docx
These templates use Jinja2 syntax for docxtpl
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_template_header(doc, titulo):
    """Add header to template"""
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(titulo)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_template_field(doc, label, placeholder):
    """Add a labeled field with Jinja2 placeholder"""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(f"{{{{ {placeholder} }}}}")

def create_familia_alimentos_template():
    """Create template for family law - alimony initial petition"""
    doc = Document()

    # Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Header info
    add_template_field(doc, "EXMO. SR. DR. JUIZ DE DIREITO DA", "foro")
    add_template_field(doc, "VARA", "vara")
    add_template_field(doc, "COMARCA", "comarca")
    doc.add_paragraph()

    # Parties
    p = doc.add_paragraph()
    p.add_run("{{ autor.nome }}").bold = True
    p.add_run(", {{ autor.qualificacao }}, ")
    p.add_run("brasileiro(a), {{ autor.estado_civil }}, {{ autor.profissao }}, ")
    p.add_run("portador(a) do CPF nº {{ autor.cpf }}, residente e domiciliado(a) ")
    p.add_run("na {{ autor.endereco }}, por seu advogado que esta subscreve ")
    p.add_run("(OAB/{{ advogado.oab }}), com escritório profissional na ")
    p.add_run("{{ advogado.endereco }}, vem, respeitosamente, à presença ")
    p.add_run("de Vossa Excelência, propor:")
    doc.add_paragraph()

    # Action title
    title = doc.add_paragraph()
    title_run = title.add_run("AÇÃO DE ALIMENTOS")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("em face de ").bold = False
    p.add_run("{{ reu.nome }}").bold = True
    p.add_run(", {{ reu.qualificacao }}, ")
    p.add_run("brasileiro(a), {{ reu.estado_civil }}, {{ reu.profissao }}, ")
    p.add_run("portador(a) do CPF nº {{ reu.cpf }}, residente e domiciliado(a) ")
    p.add_run("na {{ reu.endereco }}, pelos fatos e fundamentos a seguir expostos:")
    doc.add_paragraph()

    # DOS FATOS
    heading = doc.add_heading('I - DOS FATOS', level=1)
    doc.add_paragraph("{{ fatos }}")
    doc.add_paragraph()

    # DO DIREITO
    heading = doc.add_heading('II - DO DIREITO', level=1)
    doc.add_paragraph("{{ direito }}")
    doc.add_paragraph()

    # DA FUNDAMENTAÇÃO LEGAL
    heading = doc.add_heading('III - DA FUNDAMENTAÇÃO LEGAL E JURISPRUDENCIAL', level=1)

    # Legal citations loop
    doc.add_paragraph("{% for citacao in fundamentos_legais %}")
    p = doc.add_paragraph()
    p.add_run("• {{ citacao.titulo }}: ").bold = True
    p.add_run("{{ citacao.texto }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    # Jurisprudence citations loop
    doc.add_paragraph("{% for citacao in jurisprudencia %}")
    p = doc.add_paragraph()
    p.add_run("{{ citacao.tribunal }} - {{ citacao.numero }}:").bold = True
    doc.add_paragraph("\"{{ citacao.ementa }}\"")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    # DOS PEDIDOS
    heading = doc.add_heading('IV - DOS PEDIDOS', level=1)
    doc.add_paragraph("Diante do exposto, requer a Vossa Excelência:")
    doc.add_paragraph()

    doc.add_paragraph("{% for pedido in pedidos %}")
    doc.add_paragraph("{{ loop.index }}) {{ pedido }};")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    # DO VALOR DA CAUSA
    heading = doc.add_heading('V - DO VALOR DA CAUSA', level=1)
    doc.add_paragraph("Dá-se à causa o valor de R$ {{ valor_causa }}")
    doc.add_paragraph()

    # Closing
    doc.add_paragraph("Termos em que,")
    doc.add_paragraph("Pede deferimento.")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("{{ comarca }}, {{ data }}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("_" * 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("{{ advogado.nome }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("OAB/{{ advogado.oab }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

def create_bancario_pix_template():
    """Create template for banking - PIX fraud initial petition"""
    doc = Document()

    # Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Header
    add_template_field(doc, "EXMO. SR. DR. JUIZ DE DIREITO DO", "foro")
    add_template_field(doc, "JUIZADO ESPECIAL CÍVEL", "juizado")
    add_template_field(doc, "COMARCA", "comarca")
    doc.add_paragraph()

    # Parties
    p = doc.add_paragraph()
    p.add_run("{{ autor.nome }}").bold = True
    p.add_run(", {{ autor.qualificacao }}, inscrito(a) no CPF sob o nº {{ autor.cpf }}, ")
    p.add_run("residente e domiciliado(a) em {{ autor.endereco }}, ")
    p.add_run("por intermédio de seu advogado signatário (OAB/{{ advogado.oab }}), ")
    p.add_run("vem, respeitosamente, perante Vossa Excelência, propor a presente:")
    doc.add_paragraph()

    # Action title
    title = doc.add_paragraph()
    title_run = title.add_run("AÇÃO DE INDENIZAÇÃO POR DANOS MATERIAIS E MORAIS\nFRAUDE PIX - FALHA NA PRESTAÇÃO DE SERVIÇO")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("em face de ").bold = False
    p.add_run("{{ reu.nome }}").bold = True
    p.add_run(", pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {{ reu.cnpj }}, ")
    p.add_run("com sede em {{ reu.endereco }}, ")
    p.add_run("pelos fatos e fundamentos jurídicos a seguir aduzidos:")
    doc.add_paragraph()

    # DOS FATOS
    heading = doc.add_heading('I - DOS FATOS', level=1)
    doc.add_paragraph("{{ fatos }}")

    p = doc.add_paragraph()
    p.add_run("Valor da transferência fraudulenta: ").bold = True
    p.add_run("R$ {{ valor_fraude }}")

    p = doc.add_paragraph()
    p.add_run("Data da ocorrência: ").bold = True
    p.add_run("{{ data_fraude }}")
    doc.add_paragraph()

    # DO DIREITO
    heading = doc.add_heading('II - DO DIREITO', level=1)

    subheading = doc.add_heading('2.1 - Da Responsabilidade Objetiva (CDC art. 14)', level=2)
    doc.add_paragraph("{{ responsabilidade_objetiva }}")

    subheading = doc.add_heading('2.2 - Da Inversão do Ônus da Prova (CDC art. 6º, VIII)', level=2)
    doc.add_paragraph("{{ inversao_onus }}")

    subheading = doc.add_heading('2.3 - Dos Danos Materiais', level=2)
    doc.add_paragraph("{{ danos_materiais }}")

    subheading = doc.add_heading('2.4 - Dos Danos Morais', level=2)
    doc.add_paragraph("{{ danos_morais }}")
    doc.add_paragraph()

    # FUNDAMENTAÇÃO
    heading = doc.add_heading('III - DA FUNDAMENTAÇÃO LEGAL E JURISPRUDENCIAL', level=1)

    doc.add_paragraph("{% for citacao in fundamentos_legais %}")
    p = doc.add_paragraph()
    p.add_run("{{ citacao.titulo }}").bold = True
    doc.add_paragraph("{{ citacao.texto }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph("{% for jur in jurisprudencia %}")
    p = doc.add_paragraph()
    p.add_run("{{ jur.orgao }} - {{ jur.tipo }} {{ jur.numero }}:").bold = True
    doc.add_paragraph("{{ jur.ementa }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    # DOS PEDIDOS
    heading = doc.add_heading('IV - DOS PEDIDOS', level=1)
    doc.add_paragraph("Diante de todo o exposto, requer-se a Vossa Excelência:")
    doc.add_paragraph()

    doc.add_paragraph("{% for pedido in pedidos %}")
    doc.add_paragraph("{{ loop.index }}. {{ pedido }};")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    # VALOR DA CAUSA
    heading = doc.add_heading('V - DO VALOR DA CAUSA', level=1)
    doc.add_paragraph("Para os devidos fins legais, atribui-se à presente demanda o valor de R$ {{ valor_causa }}.")
    doc.add_paragraph()

    # PROVAS
    heading = doc.add_heading('VI - DAS PROVAS', level=1)
    doc.add_paragraph("Requer a produção de todos os meios de prova em direito admitidos, especialmente:")
    doc.add_paragraph("• Prova documental (extratos bancários, protocolos de atendimento);")
    doc.add_paragraph("• Prova testemunhal, se necessário;")
    doc.add_paragraph("• Depoimento pessoal do representante legal da Ré.")
    doc.add_paragraph()

    # Closing
    doc.add_paragraph("Nestes termos,")
    doc.add_paragraph("Pede deferimento.")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("{{ comarca }}, {{ data }}.")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("_" * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("{{ advogado.nome }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("OAB/{{ advogado.oab }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

def create_plano_saude_template():
    """Create template for health insurance denial initial petition"""
    doc = Document()

    # Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Header
    add_template_field(doc, "EXMO. SR. DR. JUIZ DE DIREITO DA", "foro")
    add_template_field(doc, "VARA CÍVEL", "vara")
    add_template_field(doc, "COMARCA", "comarca")
    doc.add_paragraph()

    # Parties
    p = doc.add_paragraph()
    p.add_run("{{ autor.nome }}").bold = True
    p.add_run(", {{ autor.qualificacao }}, portador(a) do CPF nº {{ autor.cpf }} e RG nº {{ autor.rg }}, ")
    p.add_run("residente e domiciliado(a) em {{ autor.endereco }}, ")
    p.add_run("por intermédio de seu advogado (OAB/{{ advogado.oab }}), ")
    p.add_run("vem, respeitosamente, à presença de Vossa Excelência, propor:")
    doc.add_paragraph()

    # Action title
    title = doc.add_paragraph()
    title_run = title.add_run("AÇÃO DE OBRIGAÇÃO DE FAZER C/C INDENIZAÇÃO\nNEGATIVA DE COBERTURA - PLANO DE SAÚDE")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("em face de ").bold = False
    p.add_run("{{ reu.nome }}").bold = True
    p.add_run(", pessoa jurídica de direito privado, inscrita no CNPJ {{ reu.cnpj }}, ")
    p.add_run("com sede em {{ reu.endereco }}, ")
    p.add_run("operadora de plano de saúde registrada na ANS sob nº {{ reu.ans }}, ")
    p.add_run("pelos fundamentos fáticos e jurídicos a seguir expostos:")
    doc.add_paragraph()

    # DOS FATOS
    heading = doc.add_heading('I - DA SÍNTESE FÁTICA', level=1)

    p = doc.add_paragraph()
    p.add_run("Contrato de plano de saúde nº: ").bold = True
    p.add_run("{{ contrato_numero }}")

    p = doc.add_paragraph()
    p.add_run("Vigência desde: ").bold = True
    p.add_run("{{ contrato_inicio }}")

    p = doc.add_paragraph()
    p.add_run("Carência cumprida: ").bold = True
    p.add_run("{{ carencia_cumprida }}")
    doc.add_paragraph()

    doc.add_paragraph("{{ fatos }}")

    p = doc.add_paragraph()
    p.add_run("Procedimento negado: ").bold = True
    p.add_run("{{ procedimento_negado }}")

    p = doc.add_paragraph()
    p.add_run("CID: ").bold = True
    p.add_run("{{ cid }}")

    p = doc.add_paragraph()
    p.add_run("Justificativa da negativa: ").bold = True
    p.add_run("{{ justificativa_negativa }}")
    doc.add_paragraph()

    # DO DIREITO
    heading = doc.add_heading('II - DO DIREITO', level=1)

    subheading = doc.add_heading('2.1 - Da Abusividade da Negativa de Cobertura', level=2)
    doc.add_paragraph("{{ abusividade }}")

    subheading = doc.add_heading('2.2 - Do Rol da ANS como Referência Mínima', level=2)
    doc.add_paragraph("{{ rol_ans }}")

    subheading = doc.add_heading('2.3 - Da Prevalência da Prescrição Médica', level=2)
    doc.add_paragraph("{{ prescricao_medica }}")

    subheading = doc.add_heading('2.4 - Dos Danos Morais', level=2)
    doc.add_paragraph("{{ danos_morais }}")
    doc.add_paragraph()

    # FUNDAMENTAÇÃO
    heading = doc.add_heading('III - DA FUNDAMENTAÇÃO LEGAL E JURISPRUDENCIAL', level=1)

    subheading = doc.add_heading('3.1 - Legislação Aplicável', level=2)
    doc.add_paragraph("{% for lei in fundamentos_legais %}")
    p = doc.add_paragraph()
    p.add_run("• {{ lei.titulo }}: ").bold = True
    doc.add_paragraph("{{ lei.texto }}")
    doc.add_paragraph("{% endfor %}")

    subheading = doc.add_heading('3.2 - Regulação ANS', level=2)
    doc.add_paragraph("{% for regulacao in regulacao_ans %}")
    p = doc.add_paragraph()
    p.add_run("• {{ regulacao.titulo }}: ").bold = True
    doc.add_paragraph("{{ regulacao.texto }}")
    doc.add_paragraph("{% endfor %}")

    subheading = doc.add_heading('3.3 - Jurisprudência', level=2)
    doc.add_paragraph("{% for jur in jurisprudencia %}")
    p = doc.add_paragraph()
    p.add_run("{{ jur.tribunal }} - {{ jur.tipo }} {{ jur.numero }}").bold = True
    doc.add_paragraph("{{ jur.ementa }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    # TUTELA DE URGÊNCIA
    heading = doc.add_heading('IV - DO PEDIDO DE TUTELA DE URGÊNCIA', level=1)
    doc.add_paragraph("{{ tutela_urgencia }}")
    doc.add_paragraph()

    # DOS PEDIDOS
    heading = doc.add_heading('V - DOS PEDIDOS', level=1)
    doc.add_paragraph("Diante de todo o exposto, requer a Vossa Excelência:")
    doc.add_paragraph()

    doc.add_paragraph("{% for pedido in pedidos %}")
    doc.add_paragraph("{{ loop.index }}. {{ pedido }};")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    # VALOR DA CAUSA
    heading = doc.add_heading('VI - DO VALOR DA CAUSA', level=1)
    doc.add_paragraph("Atribui-se à causa o valor de R$ {{ valor_causa }}.")
    doc.add_paragraph()

    # PROVAS
    heading = doc.add_heading('VII - DAS PROVAS', level=1)
    doc.add_paragraph("Protesta pela produção de todos os meios de prova em direito admitidos, em especial:")
    doc.add_paragraph("• Documental (relatórios médicos, negativa escrita, contrato);")
    doc.add_paragraph("• Pericial, se necessário;")
    doc.add_paragraph("• Depoimento pessoal do representante legal da operadora.")
    doc.add_paragraph()

    # Closing
    doc.add_paragraph("Termos em que,")
    doc.add_paragraph("Pede deferimento.")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("{{ comarca }}, {{ data }}.")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("_" * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("{{ advogado.nome }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("OAB/{{ advogado.oab }}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

def main():
    """Generate all three templates"""
    templates_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
    os.makedirs(templates_dir, exist_ok=True)

    print("Generating DOCX templates...")

    # 1. Family law - Alimony
    print("1/3 Creating inicial_familia_alimentos.docx...")
    doc1 = create_familia_alimentos_template()
    doc1.save(os.path.join(templates_dir, 'inicial_familia_alimentos.docx'))
    print("✓ inicial_familia_alimentos.docx created")

    # 2. Banking - PIX fraud
    print("2/3 Creating inicial_bancario_pix.docx...")
    doc2 = create_bancario_pix_template()
    doc2.save(os.path.join(templates_dir, 'inicial_bancario_pix.docx'))
    print("✓ inicial_bancario_pix.docx created")

    # 3. Health insurance
    print("3/3 Creating inicial_plano_saude.docx...")
    doc3 = create_plano_saude_template()
    doc3.save(os.path.join(templates_dir, 'inicial_plano_saude.docx'))
    print("✓ inicial_plano_saude.docx created")

    print("\n✅ All templates created successfully!")
    print(f"Location: {templates_dir}")

if __name__ == "__main__":
    main()
