"""
PDF and document generation services
"""
import os
from datetime import datetime
from typing import Dict, List
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docxtpl import DocxTemplate
from io import BytesIO

# Paths
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")
DOCS_DIR = os.path.join(TEMPLATES_DIR, "docs")
REPORTS_DIR = os.getenv("REPORTS_DIR", "/reports")

# Ensure directories exist
os.makedirs(REPORTS_DIR, exist_ok=True)

# Jinja2 environment for HTML templates
jinja_env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))


def generate_pdf_report(case, data_atualizacao: str) -> str:
    """
    Generate premium PDF report for a case

    Args:
        case: Case model instance
        data_atualizacao: Date string for corpus update

    Returns:
        Path to generated PDF file
    """
    # Load HTML template
    template = jinja_env.get_template("report.html")

    # Prepare data
    context = {
        "case_id": case.id,
        "data_criacao": case.created_at.strftime("%d/%m/%Y"),
        "data_atualizacao": data_atualizacao,
        "area": case.area or "Não especificada",
        "sub_area": case.sub_area or "",
        "tipificacao": case.typification or "",
        "estrategias": case.strategies or "",
        "riscos": case.risks or "",
        "probabilidade": case.probability.value if case.probability else "média",
        "custos": case.cost_estimate or "",
        "prazos": case.timeline_estimate or "",
        "checklist": case.checklist or [],
        "rascunho": case.draft_petition or "",
        "citacoes": case.citations or [],
    }

    # Render HTML
    html_content = template.render(**context)

    # Generate PDF
    pdf_filename = f"relatorio_{case.id}_{int(datetime.utcnow().timestamp())}.pdf"
    pdf_path = os.path.join(REPORTS_DIR, pdf_filename)

    HTML(string=html_content).write_pdf(pdf_path)

    return pdf_path


def generate_document(
    tipo_peca: str,
    area: str,
    metadata: Dict,
    blocks: Dict,
    citacoes: List,
    format: str = "docx"
) -> str:
    """
    Generate legal document (petition, contestation, appeal) from template

    Args:
        tipo_peca: Type of document (inicial, contestacao, recurso)
        area: Legal area (familia, consumidor, bancario, saude)
        metadata: Case and party metadata
        blocks: LLM-generated text blocks
        citacoes: List of citations
        format: Output format (docx or pdf)

    Returns:
        Path to generated document
    """
    # Select template based on area and tipo_peca
    template_name = f"modelo_{tipo_peca}_{area}.docx"
    template_path = os.path.join(DOCS_DIR, template_name)

    # Fallback to generic template if specific not found
    if not os.path.exists(template_path):
        template_name = f"modelo_{tipo_peca}_generico.docx"
        template_path = os.path.join(DOCS_DIR, template_name)

    # If still not found, create simple document
    if not os.path.exists(template_path):
        return generate_simple_document(metadata, blocks, citacoes, format)

    # Load template
    doc = DocxTemplate(template_path)

    # Prepare context for template
    from services.citations import CitationManager
    citation_mgr = CitationManager()

    # Add citations
    citacoes_numeradas = []
    for cit in citacoes:
        numero = citation_mgr.add_citation(cit.model_dump() if hasattr(cit, 'model_dump') else cit)
        citacoes_numeradas.append({
            "numero": numero,
            "titulo": cit.get("titulo", "") if isinstance(cit, dict) else cit.titulo,
            "texto": cit.get("texto", "") if isinstance(cit, dict) else cit.texto,
            "orgao": cit.get("orgao", "") if isinstance(cit, dict) else cit.orgao,
            "tribunal": cit.get("tribunal", "") if isinstance(cit, dict) else cit.tribunal,
        })

    # Build context
    context = {
        # Metadata
        "foro": metadata.get("foro", ""),
        "vara": metadata.get("vara", ""),
        "tipo_acao": metadata.get("tipo_acao", tipo_peca.upper()),
        "autor": {
            "nome": metadata.get("autor_nome", ""),
            "qualificacao": metadata.get("autor_qualificacao", "")
        },
        "reu": {
            "nome": metadata.get("reu_nome", ""),
            "qualificacao": metadata.get("reu_qualificacao", "")
        },
        "advogado": {
            "nome": metadata.get("advogado_nome", ""),
            "oab": metadata.get("advogado_oab", "")
        },
        "valor_causa": metadata.get("valor_causa", "0,00"),

        # Generated blocks
        "fatos_texto": blocks.get("fatos_detalhados", metadata.get("fatos_resumo", "")),
        "fundamentacao": blocks.get("fundamentacao_juridica", ""),
        "pedidos": blocks.get("pedidos_elaborados", []),

        # Citations
        "fundamentos": citacoes_numeradas,
        "jurisprudencia": [c for c in citacoes_numeradas if c.get("tribunal")],

        # Date
        "cidade": "São Paulo",  # TODO: get from metadata
        "data_extenso": datetime.now().strftime("%d de %B de %Y")
    }

    # Render template
    doc.render(context)

    # Save document
    doc_filename = f"peca_{tipo_peca}_{int(datetime.utcnow().timestamp())}.docx"
    doc_path = os.path.join(REPORTS_DIR, doc_filename)

    doc.save(doc_path)

    # Convert to PDF if requested
    if format == "pdf":
        pdf_path = convert_docx_to_pdf(doc_path)
        return pdf_path

    return doc_path


def generate_simple_document(metadata: Dict, blocks: Dict, citacoes: List, format: str) -> str:
    """Generate a simple document without template"""
    doc = Document()

    # Title
    doc.add_heading(f'{metadata.get("tipo_acao", "PETIÇÃO INICIAL")}', 0)

    # Header
    doc.add_paragraph(f'Excelentíssimo Senhor Doutor Juiz de Direito da {metadata.get("vara", "")} de {metadata.get("foro", "")}')
    doc.add_paragraph()

    # Parties
    doc.add_paragraph(
        f'{metadata.get("autor_nome", "")}, {metadata.get("autor_qualificacao", "")}, '
        f'por seu advogado (OAB {metadata.get("advogado_oab", "")}), propõe'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f'{metadata.get("tipo_acao", "AÇÃO")} em face de {metadata.get("reu_nome", "")}, '
        f'{metadata.get("reu_qualificacao", "")}, pelos fatos e fundamentos a seguir expostos:'
    )
    doc.add_paragraph()

    # Facts
    doc.add_heading('I. DOS FATOS', 1)
    doc.add_paragraph(blocks.get("fatos_detalhados", ""))
    doc.add_paragraph()

    # Legal basis
    doc.add_heading('II. DO DIREITO', 1)
    doc.add_paragraph(blocks.get("fundamentacao_juridica", ""))
    doc.add_paragraph()

    # Requests
    doc.add_heading('III. DOS PEDIDOS', 1)
    pedidos = blocks.get("pedidos_elaborados", [])
    for i, pedido in enumerate(pedidos, 1):
        doc.add_paragraph(f'{i}. {pedido}')
    doc.add_paragraph()

    # Closing
    doc.add_paragraph(f'Protesta por provas em direito admitidas.')
    doc.add_paragraph(f'Dá-se à causa o valor de R$ {metadata.get("valor_causa", "0,00")}.')
    doc.add_paragraph()
    doc.add_paragraph(f'{metadata.get("cidade", "São Paulo")}, {datetime.now().strftime("%d de %B de %Y")}.')
    doc.add_paragraph()
    doc.add_paragraph(f'_________________________________')
    doc.add_paragraph(f'{metadata.get("advogado_nome", "")} - OAB {metadata.get("advogado_oab", "")}')

    # Save
    doc_filename = f"peca_{int(datetime.utcnow().timestamp())}.docx"
    doc_path = os.path.join(REPORTS_DIR, doc_filename)
    doc.save(doc_path)

    if format == "pdf":
        return convert_docx_to_pdf(doc_path)

    return doc_path


def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert DOCX to PDF

    Note: This is a stub. In production, use LibreOffice, pandoc, or docx2pdf
    """
    # For MVP, we'll just return the DOCX path
    # In production, use: docx2pdf, pypandoc, or libreoffice --convert-to pdf

    pdf_path = docx_path.replace(".docx", ".pdf")

    # Stub: just copy the file (in production, do actual conversion)
    # from docx2pdf import convert
    # convert(docx_path, pdf_path)

    # For now, return DOCX path as fallback
    return docx_path
