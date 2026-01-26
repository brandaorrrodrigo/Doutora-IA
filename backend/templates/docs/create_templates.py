"""
Script to create DOCX templates for legal documents
Run this once to generate the template files
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_inicial_familia_template():
    """Create template for initial family law petition"""
    doc = Document()

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run('{{ foro }}\n{{ vara }}')

    doc.add_paragraph()

    # Parties
    p = doc.add_paragraph()
    p.add_run('{{ autor.nome }}').bold = True
    p.add_run(', {{ autor.qualificacao }}, por seu advogado (OAB {{ advogado.oab }}), propõe')

    doc.add_paragraph()

    # Title
    title = doc.add_heading('{{ tipo_acao|upper }}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph('em face de ')
    p.add_run('{{ reu.nome }}').bold = True
    p.add_run(', {{ reu.qualificacao }}, pelos fatos e fundamentos a seguir expostos:')

    doc.add_paragraph()

    # I. DOS FATOS
    doc.add_heading('I. DOS FATOS', 1)
    doc.add_paragraph('{{ fatos_texto }}')
    doc.add_paragraph()

    # II. DO DIREITO
    doc.add_heading('II. DO DIREITO', 1)
    doc.add_paragraph('{% for cit in fundamentos %}')
    doc.add_paragraph('[{{ loop.index }}] {{ cit.titulo }}')
    doc.add_paragraph('{{ cit.texto }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # III. DA JURISPRUDÊNCIA
    doc.add_heading('III. DA JURISPRUDÊNCIA', 1)
    doc.add_paragraph('{% for j in jurisprudencia %}')
    doc.add_paragraph('[{{ loop.index }}] {{ j.tribunal }} – {{ j.classe }} nº {{ j.numero }}')
    doc.add_paragraph('Ementa: {{ j.ementa_curta }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # IV. DOS PEDIDOS
    doc.add_heading('IV. DOS PEDIDOS', 1)
    doc.add_paragraph('Diante do exposto, requer-se a Vossa Excelência:')
    doc.add_paragraph()
    doc.add_paragraph('{% for p in pedidos %}')
    doc.add_paragraph('{{ loop.index }}. {{ p }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # Closing
    doc.add_paragraph('Protesta por provas em direito admitidas, em especial documental, testemunhal e pericial.')
    doc.add_paragraph()
    doc.add_paragraph('Dá-se à causa o valor de R$ {{ valor_causa }}.')
    doc.add_paragraph()
    doc.add_paragraph('{{ cidade }}, {{ data_extenso }}.')
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph('_________________________________')
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig2 = doc.add_paragraph('{{ advogado.nome }}')
    sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig3 = doc.add_paragraph('OAB {{ advogado.oab }}')
    sig3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def create_inicial_pix_template():
    """Create template for PIX fraud/banking petition"""
    doc = Document()

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run('{{ foro }}\n{{ vara }}')

    doc.add_paragraph()

    # Parties
    p = doc.add_paragraph()
    p.add_run('{{ autor.nome }}').bold = True
    p.add_run(', {{ autor.qualificacao }}, por seu advogado (OAB {{ advogado.oab }}), propõe')

    doc.add_paragraph()

    # Title
    title = doc.add_heading('AÇÃO DE RESSARCIMENTO DE DANOS C/C PEDIDO DE TUTELA ANTECIPADA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph('em face de ')
    p.add_run('{{ reu.nome }}').bold = True
    p.add_run(', {{ reu.qualificacao }}, pelos fatos e fundamentos a seguir expostos:')

    doc.add_paragraph()

    # I. DOS FATOS
    doc.add_heading('I. DOS FATOS', 1)
    doc.add_paragraph('{{ fatos_texto }}')
    doc.add_paragraph()

    # II. DA RESPONSABILIDADE DO BANCO
    doc.add_heading('II. DA RESPONSABILIDADE OBJETIVA DA INSTITUIÇÃO FINANCEIRA', 1)
    doc.add_paragraph('{% for cit in fundamentos %}')
    doc.add_paragraph('[{{ loop.index }}] {{ cit.titulo }}')
    doc.add_paragraph('{{ cit.texto }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # III. DA JURISPRUDÊNCIA
    doc.add_heading('III. DA JURISPRUDÊNCIA PACÍFICA', 1)
    doc.add_paragraph('{% for j in jurisprudencia %}')
    doc.add_paragraph('[{{ loop.index }}] {{ j.tribunal }} – {{ j.classe }} nº {{ j.numero }}')
    doc.add_paragraph('Ementa: {{ j.ementa_curta }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # IV. DA TUTELA ANTECIPADA
    doc.add_heading('IV. DA TUTELA ANTECIPADA', 1)
    doc.add_paragraph('Presentes os requisitos do art. 300 do CPC, requer-se a concessão de tutela antecipada para determinar o estorno imediato do valor indevidamente transferido, sob pena de multa diária.')
    doc.add_paragraph()

    # V. DOS PEDIDOS
    doc.add_heading('V. DOS PEDIDOS', 1)
    doc.add_paragraph('Diante do exposto, requer-se a Vossa Excelência:')
    doc.add_paragraph()
    doc.add_paragraph('{% for p in pedidos %}')
    doc.add_paragraph('{{ loop.index }}. {{ p }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # Closing
    doc.add_paragraph('Protesta por provas em direito admitidas.')
    doc.add_paragraph()
    doc.add_paragraph('Dá-se à causa o valor de R$ {{ valor_causa }}.')
    doc.add_paragraph()
    doc.add_paragraph('{{ cidade }}, {{ data_extenso }}.')
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph('_________________________________')
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig2 = doc.add_paragraph('{{ advogado.nome }}')
    sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig3 = doc.add_paragraph('OAB {{ advogado.oab }}')
    sig3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def create_inicial_plano_saude_template():
    """Create template for health insurance petition"""
    doc = Document()

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run('{{ foro }}\n{{ vara }}')

    doc.add_paragraph()

    # Parties
    p = doc.add_paragraph()
    p.add_run('{{ autor.nome }}').bold = True
    p.add_run(', {{ autor.qualificacao }}, por seu advogado (OAB {{ advogado.oab }}), propõe')

    doc.add_paragraph()

    # Title
    title = doc.add_heading('AÇÃO DE OBRIGAÇÃO DE FAZER C/C PEDIDO LIMINAR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph('em face de ')
    p.add_run('{{ reu.nome }}').bold = True
    p.add_run(', {{ reu.qualificacao }}, pelos fatos e fundamentos a seguir expostos:')

    doc.add_paragraph()

    # I. DOS FATOS
    doc.add_heading('I. DOS FATOS', 1)
    doc.add_paragraph('{{ fatos_texto }}')
    doc.add_paragraph()

    # II. DO DIREITO À SAÚDE
    doc.add_heading('II. DO DIREITO À COBERTURA E DO ROL DA ANS', 1)
    doc.add_paragraph('{% for cit in fundamentos %}')
    doc.add_paragraph('[{{ loop.index }}] {{ cit.titulo }}')
    doc.add_paragraph('{{ cit.texto }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # III. DA JURISPRUDÊNCIA
    doc.add_heading('III. DA JURISPRUDÊNCIA', 1)
    doc.add_paragraph('{% for j in jurisprudencia %}')
    doc.add_paragraph('[{{ loop.index }}] {{ j.tribunal }} – {{ j.classe }} nº {{ j.numero }}')
    doc.add_paragraph('Ementa: {{ j.ementa_curta }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # IV. DO PEDIDO LIMINAR
    doc.add_heading('IV. DO PEDIDO LIMINAR', 1)
    doc.add_paragraph('Diante da urgência e risco de dano irreparável à saúde do autor, requer-se a concessão de liminar inaudita altera parte para determinar a realização imediata do procedimento/fornecimento do medicamento, sob pena de multa diária.')
    doc.add_paragraph()

    # V. DOS PEDIDOS
    doc.add_heading('V. DOS PEDIDOS', 1)
    doc.add_paragraph('Diante do exposto, requer-se a Vossa Excelência:')
    doc.add_paragraph()
    doc.add_paragraph('{% for p in pedidos %}')
    doc.add_paragraph('{{ loop.index }}. {{ p }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph()

    # Closing
    doc.add_paragraph('Protesta por provas em direito admitidas.')
    doc.add_paragraph()
    doc.add_paragraph('Dá-se à causa o valor de R$ {{ valor_causa }}.')
    doc.add_paragraph()
    doc.add_paragraph('{{ cidade }}, {{ data_extenso }}.')
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph('_________________________________')
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig2 = doc.add_paragraph('{{ advogado.nome }}')
    sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig3 = doc.add_paragraph('OAB {{ advogado.oab }}')
    sig3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


if __name__ == "__main__":
    # Get script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Create templates
    print("Creating template: modelo_inicial_familia.docx")
    doc1 = create_inicial_familia_template()
    doc1.save(os.path.join(script_dir, "modelo_inicial_familia.docx"))

    print("Creating template: modelo_inicial_pix.docx")
    doc2 = create_inicial_pix_template()
    doc2.save(os.path.join(script_dir, "modelo_inicial_pix.docx"))

    print("Creating template: modelo_inicial_plano_saude.docx")
    doc3 = create_inicial_plano_saude_template()
    doc3.save(os.path.join(script_dir, "modelo_inicial_plano_saude.docx"))

    print("✓ All templates created successfully!")
